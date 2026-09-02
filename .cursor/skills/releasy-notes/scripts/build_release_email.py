#!/usr/bin/env python3
"""
Build a short, business-facing "Xeelo - Release Notes" email-style Word document.

This script only handles the mechanical part: turning already-written content
(title/version/date + lists of plain-English paragraphs) into a .docx file
with the exact styling of the reference template. It does NOT read the raw
DevOps ticket list or write the paragraph text - that judgment call (deciding
what's a Feature vs a Bug, translating/rewriting raw ticket text into clean
business prose, writing the summary) belongs to whoever is authoring the
content, per SKILL.md.

Usage:
    python build_release_email.py content.json output.docx

content.json schema:
{
  "version": "Labe-07.013",
  "release_date": "18 August 2026",
  "notice": {
    "title": "Optional callout heading.",
    "paragraphs": ["Optional highlighted box, rendered at the top in a red frame."]
  },
  "features": ["Paragraph 1 text.", "Paragraph 2 text.", ...],
  "bugs": ["Paragraph 1 text.", ...],
  "summary": ["Paragraph 1 text.", "Optional paragraph 2 text."]
}

Every string in "features"/"bugs" is rendered as one paragraph, formatted as
"Title. Explanation..." in a single run (no bold inside the paragraph body -
only the section headings and the Version:/Release date: labels are bold).
"summary" may be one paragraph or a short list of paragraphs.
"""
import sys
import json
from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
TITLE_COLOR = "1F3864"
HEADING_COLOR = "1F3864"
BODY_COLOR = "222222"
NOTICE_BORDER = "C00000"
NOTICE_SHADING = "FDECEC"
NOTICE_TITLE_COLOR = "9B1B1B"

# Schema order for hand-built OOXML. Inserting out of order makes Word
# treat the file as needing repair and silently drop the formatting.
TBLPR_ORDER = (
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
    "tblStyleRowBandSize", "tblStyleColBandSize",
    "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar",
    "tblLook", "tblCaption", "tblDescription",
)
TCPR_ORDER = (
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge",
    "tcBorders", "shd", "noWrap", "tcMar",
    "textDirection", "tcFitText", "vAlign", "hideMark", "headers",
)


def set_run_font(run, size_pt, color_hex, bold=False):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color_hex)


def add_plain_paragraph(doc, text, size_pt, color_hex, bold=False,
                         before_pt=0, after_pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)
    r = p.add_run(text)
    set_run_font(r, size_pt, color_hex, bold)
    return p


def add_label_value_paragraph(doc, label, value, size_pt=11,
                               color_hex=BODY_COLOR, before_pt=0, after_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)
    r1 = p.add_run(label)
    set_run_font(r1, size_pt, color_hex, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size_pt, color_hex, bold=False)
    return p


def add_heading(doc, text):
    return add_plain_paragraph(doc, text, 13, HEADING_COLOR, bold=True,
                                before_pt=16, after_pt=6)


def add_item_paragraph(doc, text):
    return add_plain_paragraph(doc, text, 11, BODY_COLOR, bold=False,
                                before_pt=0, after_pt=8)


def _local_tag(element):
    tag = element.tag
    return tag.split("}")[-1] if "}" in tag else tag


def _insert_in_schema_order(parent, child, order):
    """Insert child before the first existing sibling that comes later in order."""
    name = _local_tag(child)
    try:
        idx = order.index(name)
    except ValueError:
        parent.append(child)
        return
    later = set(order[idx + 1:])
    for existing in list(parent):
        if _local_tag(existing) in later:
            existing.addprevious(child)
            return
    parent.append(child)


def _set_borders(parent, order, color_hex, sz="16"):
    tag = "w:tblBorders" if "tblBorders" in order else "w:tcBorders"
    borders = parent.find(qn(tag))
    if borders is not None:
        parent.remove(borders)
    borders = OxmlElement(tag)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    for edge in ("insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    _insert_in_schema_order(parent, borders, order)


def _set_shading(parent, order, fill_hex):
    existing = parent.find(qn("w:shd"))
    if existing is not None:
        parent.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    _insert_in_schema_order(parent, shd, order)


def _set_tbl_width(tbl_pr, twips):
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(twips))
    tbl_w.set(qn("w:type"), "dxa")
    _insert_in_schema_order(tbl_pr, tbl_w, TBLPR_ORDER)


def _set_cell_margins(tc_pr, twips=120):
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(twips))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    _insert_in_schema_order(tc_pr, tc_mar, TCPR_ORDER)


def _clear_cell_paragraphs(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)


def add_notice_box(doc, title, paragraphs):
    """Single-cell table with a red frame. Used for migration / cutover notices."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    # Content width: 12240 - 1260 - 1260 = 9720 twips
    _set_tbl_width(tbl_pr, 9720)
    _set_borders(tbl_pr, TBLPR_ORDER, NOTICE_BORDER)
    _set_shading(tbl_pr, TBLPR_ORDER, NOTICE_SHADING)

    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    _set_borders(tc_pr, TCPR_ORDER, NOTICE_BORDER)
    _set_shading(tc_pr, TCPR_ORDER, NOTICE_SHADING)
    _set_cell_margins(tc_pr, 140)
    _clear_cell_paragraphs(cell)

    if title:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        set_run_font(r, 11, NOTICE_TITLE_COLOR, bold=True)

    for i, text in enumerate(paragraphs):
        if not text:
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if i == len(paragraphs) - 1 else 6)
        r = p.add_run(text)
        set_run_font(r, 11, BODY_COLOR, bold=False)

    # Breathing room after the box (tables themselves have no space_after).
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    return table


def set_page(doc):
    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Twips(1080)
    section.bottom_margin = Twips(1080)
    section.left_margin = Twips(1260)
    section.right_margin = Twips(1260)
    section.header_distance = Twips(708)
    section.footer_distance = Twips(708)


def build(content, output_path):
    doc = Document()
    set_page(doc)

    title = content.get("document_title") or "Xeelo – Release Notes"
    add_plain_paragraph(doc, title, 18, TITLE_COLOR,
                         bold=True, before_pt=0, after_pt=8)
    add_label_value_paragraph(doc, "Version: ", content["version"])
    add_label_value_paragraph(doc, "Release date: ", content["release_date"])

    notice = content.get("notice")
    if notice:
        if isinstance(notice, str):
            add_notice_box(doc, title=None, paragraphs=[notice])
        else:
            paragraphs = notice.get("paragraphs")
            if not paragraphs and notice.get("body"):
                paragraphs = [notice["body"]]
            add_notice_box(doc, title=notice.get("title"),
                           paragraphs=paragraphs or [])

    features = content.get("features", [])
    if features:
        add_heading(doc, "New Features")
        for text in features:
            add_item_paragraph(doc, text)

    bugs = content.get("bugs", [])
    if bugs:
        add_heading(doc, "Bug Fixes")
        for text in bugs:
            add_item_paragraph(doc, text)

    summary = content.get("summary", [])
    if isinstance(summary, str):
        summary = [summary]
    if summary:
        add_heading(doc, "Summary")
        for text in summary:
            add_item_paragraph(doc, text)

    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python build_release_email.py content.json output.docx",
              file=sys.stderr)
        sys.exit(1)
    content_path, output_path = sys.argv[1], sys.argv[2]
    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)
    for required in ("version", "release_date"):
        if required not in content:
            print(f"content.json is missing required field: {required}",
                  file=sys.stderr)
            sys.exit(1)
    build(content, output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
